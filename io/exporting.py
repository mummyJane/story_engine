# story_engine/exporting.py

from __future__ import annotations
from pathlib import Path
from typing import Literal
import json

from docx import Document  # pip install python-docx

from story_engine.io.project_repo import ProjectRepo
from story_engine.io.schemas import ChapterOutline


def _iter_chapters(repo: ProjectRepo):
    cfg = repo.load_project_config()
    for cid in cfg.chapter_order:
        outline_dict = repo.load_chapter_outline(cid) or {}
        title = outline_dict.get("title", cid)
        text = (
            repo.load_chapter_text(cid, kind="final")
            or repo.load_chapter_text(cid, kind="draft")
            or ""
        )
        yield cid, title, text


def export_txt(repo: ProjectRepo, out_path: Path) -> None:
    cfg = repo.load_project_config()
    with out_path.open("w", encoding="utf-8") as f:
        f.write(cfg.title + "\n\n")
        for cid, title, text in _iter_chapters(repo):
            f.write(f"{title}\n")
            f.write("=" * len(title) + "\n\n")
            f.write(text.rstrip() + "\n\n")


def export_md(repo: ProjectRepo, out_path: Path) -> None:
    cfg = repo.load_project_config()
    with out_path.open("w", encoding="utf-8") as f:
        f.write(f"# {cfg.title}\n\n")
        for cid, title, text in _iter_chapters(repo):
            f.write(f"## {title}\n\n")
            f.write(text.rstrip() + "\n\n")


def export_docx(repo: ProjectRepo, out_path: Path) -> None:
    cfg = repo.load_project_config()
    doc = Document()
    doc.add_heading(cfg.title, level=0)
    for cid, title, text in _iter_chapters(repo):
        doc.add_heading(title, level=1)
        for para in text.split("\n\n"):
            doc.add_paragraph(para.strip())
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def export_campfire(repo: ProjectRepo, out_path: Path) -> None:
    """
    Generic 'Campfire pack' JSON with chapters + bible + timeline.

    This is *not* the official Campfire format (they keep changing it),
    but it's a clean metadata pack you can transform/import as needed.
    """
    cfg = repo.load_project_config()
    bible = repo.load_bible()

    chapters = []
    for cid, title, text in _iter_chapters(repo):
        outline_dict = repo.load_chapter_outline(cid) or {}
        chapters.append(
            {
                "id": cid,
                "title": title,
                "outline": outline_dict,
                "text": text,
            }
        )

    data = {
        "project": {
            "title": cfg.title,
            "author": cfg.author,
            "tone": cfg.tone,
            "style_rules": cfg.style_rules,
            "chapter_order": cfg.chapter_order,
        },
        "chapters": chapters,
        "characters": [c.to_json_dict() if hasattr(c, "to_json_dict") else {
            "name": name,
            "bio": c.bio,
            "tags": c.tags,
            "current_state": {
                "location": c.current_state.location,
                "physical": c.current_state.physical,
                "restraints": c.current_state.restraints,
                "emotional": c.current_state.emotional,
                "knowledge": c.current_state.knowledge,
                "relationship_to": c.current_state.relationship_to,
            },
        } for name, c in bible.characters.items()],
        "locations": [
            {
                "name": name,
                "description": loc.description,
                "tags": loc.tags,
                "notes": loc.notes,
            }
            for name, loc in bible.locations.items()
        ],
        "items": [
            {
                "name": name,
                "description": it.description,
                "owner": it.owner,
                "location": it.location,
                "state": it.state,
                "tags": it.tags,
            }
            for name, it in bible.items.items()
        ],
        "timeline": [
            {
                "chapter_id": ev.chapter_id,
                "order_in_chapter": ev.order_in_chapter,
                "summary": ev.summary,
            }
            for ev in bible.timeline
        ],
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
